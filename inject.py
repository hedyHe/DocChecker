#coding=utf-8

import os
os.environ['CUDA_VISIBLE_DEVICES'] = '5,7'

import json, re
import datetime
import numpy as np
import optparse
import traceback
from collections import defaultdict
from docx import Document

from transformers import AutoModelForCausalLM, AutoTokenizer
import torch
from llm_api_vllm import Qwen

import random
random.seed(34)

class Inject(object):
	"""docstring for Inject"""
	def __init__(self):

		self.llm = 'qwen2.5_72b'

		self.llm_api = Qwen(self.llm)
		self.vllm = True
		# self.llm_api = None
		self.doc_lens = []
		self.total_sents = 0  #可以插入错误的句子

		self.doc_types = {"2024年国家免疫规划疫苗-脊髓灰质炎灭活疫苗补种集中采购项目-项目采购第七批招标公告": "招投标", "2024年国家免疫规划疫苗集中采购项目-项目采购第七批招标公告": "招投标", "中南大学湘雅医院2023年IT基础设施设备采购项目招标公告": "招投标", "中国人民大学附属中学2024-2027年度保安服务采购项目招标公告": "招投标", "中国人民银行山东省分行ACS内容缓存平台和安全云平台相关硬件设备采购项目招标公告": "招投标", "中国医学科学院北京协和医院小型机服务器采购项目招标公告": "招投标", "中国医学科学院北区物业综合管理服务采购项目招标公告": "招投标", "中国医学科学院阜外医院本部院区一号楼三层日间手术室改造工程竞争性磋商公告": "招投标", "中国医学科学院阜外医院西山园区保安服务采购项目招标公告": "招投标", "中国医学科学院阜外医院风机盘管维保采购项目招标公告": "招投标", "中国地震台网中心预警项目改扩建基础设施采购项目招标公告": "招投标", "中国地震局地质研究所园区综合物业管理服务采购项目招标公告": "招投标", "中国社会科学院大学良乡校区2024-2027年保安服务采购项目招标公告": "招投标", "中央国家机关2024年工程监理服务框架协议采购项目征集公告": "招投标", "中央美术学院消防安防控制室管理服务采购项目招标公告": "招投标", "北京大学第一医院全院印刷品采购项目招标公告": "招投标", "北京市消防救援总队应急能力建设车辆装备购置项目一标段二次招标公告": "招投标", "北京市消防救援总队应急能力建设车辆装备购置项目三标段招标公告": "招投标", "北京市消防救援总队应急能力建设车辆装备购置项目二标段招标公告": "招投标", "北京市消防救援总队应急能力建设车辆装备购置项目五标段招标公告": "招投标", "北京市消防救援总队应急能力建设车辆装备购置项目六标段招标公告": "招投标", "北京师范大学后主楼和教三楼物业管理服务采购项目招标公告": "招投标", "四川大学华西医院PACS计算存储采购项目招标公告": "招投标", "国家开放大学五棵松校区电梯更换工程采购项目（三次）资格预审公告  （工程类）  ": "招投标", "国家机关事务管理局审计室2024-2025年度审计服务框架协议采购项目征集公告": "招投标", "国家机关事务管理局西山服务局2024年7月至2027年7月食材配送采购项目招标公告": "招投标", "国家林业和草原局碳卫星海南试验站设备采购及集成项目招标公告": "招投标", "国家气象中心海洋气象综合保障二期工程远洋导航业务支撑硬件资源建设采购项目招标公告": "招投标", "国家移民管理局瑞丽遣返中心礼堂信息化建设设备采购项目招标公告": "招投标", "国家药品监督管理局医疗器械技术审评中心新址网络及多媒体融合集成采购项目招标公告": "招投标", "国家语委普通话水平测试等级证书印制项目招标公告": "招投标", "天津市地震局天津市巨灾防范工程-数据平台建设(国家中心信息化硬件系统建设)采购项目招标公告": "招投标", "天津市地震局天津市巨灾防范工程-数据平台建设(省级中心信息化建设)采购项目招标公告": "招投标", "常州市消防救援支队本级指挥中心改造采购项目招标公告": "招投标", "应急管理部2023年国家森林草原防灭火储备物资采购项目（涡轮增压型森林消防水泵）招标公告": "招投标", "教育部教育考试院数据中心计算资源采购项目招标公告": "招投标", "最高人民检察院中国-东盟检务信息交流中心综合服务系统项目采购项目招标公告": "招投标", "榆林市消防救援支队消防语音图像自组网建设和消防图像综合管理系统升级改造采购项目招标公告": "招投标", "科电大厦修缮二期(空调、消防及照明系统)项目施工招标资格预审公告": "招投标", "西安市消防救援支队培训楼会议视频系统采购项目招标公告": "招投标", "2024-2025年度新疆油田公司防腐涂料（框架）": "石油领域的招投标", "2024~2026年独山子石化中密控股密封维修框架公开招标": "石油领域的招投标", "2024年西部钻探试油测试技术研发中心建设勘察初步设计服务项目": "石油领域的招投标", "中亚天然气管道（霍尔果斯）有限公司 2024年流量计使用中测试项目（二次）": "石油领域的招投标", "中国石油天然气管道科学研究院有限公司焊接易损易耗件加工制造选商招标项目（二次）": "石油领域的招投标", "中国石油天然气股份有限公司青海销售分公司安全现状评价采购项目": "石油领域的招投标", "中国石油天然气集团有限公司2024年消防车集中采购（安全生产预防和应急救援能力建设补助资金项目）招标": "石油领域的招投标", "中国石油独山子石化公司聚烯烃二部四台往复式压缩机组采购项目公开招标二次": "石油领域的招投标", "中国石油运输有限公司甘肃化工分公司固体化工产品（庆阳石化橡塑产品补充运力）运输服务项目": "石油领域的招投标", "中国石油集团测井有限公司制造公司2024-2025年热处理等工序委外": "石油领域的招投标", "中国石油集团测井有限公司制造公司2024-2025年焊接等工序委外": "石油领域的招投标", "中油测井井下仪器在多温度、压力场中机械结构及性能仿真技术（二次）": "石油领域的招投标", "中油测井激发极化仪器数据处理方法研究及正反演软": "石油领域的招投标", "兰州石化公司化工区非抗爆控制室隐患治理项目智能巡检仪采购（二次）": "石油领域的招投标", "吉林中油宝石花石油安装有限责任公司劳务外包项目": "石油领域的招投标", "吉林石化公司转型升级项目最小回流阀采购招标项目二次": "石油领域的招投标", "吐哈油田消防车库改造项目": "石油领域的招投标", "咸阳宝石钢管钢绳有限公司2024-2025年度镀铬泵筒采购": "石油领域的招投标", "塔里木二期项目裂解气清焦阀和传输阀(变更)(变更)": "石油领域的招投标", "塔里木油田库尔勒上库高新区低碳转型130万千瓦光伏项目光伏区1标段EPC总承包": "石油领域的招投标", "抚顺石化公司烯烃厂2024年低砷磷酸框架采购（二次）": "石油领域的招投标", "新能源事业部2024-2025年度项目管理咨询服务": "石油领域的招投标", "智能化多井对比与沉积微相绘图模块研发（二次）": "石油领域的招投标", "氢氟酸代储代销": "石油领域的招投标", "测井公司吐哈分公司2024-2026年职工食堂餐饮服务业务外包项目": "石油领域的招投标", "玉门油田2024-2025炼化总厂易腐蚀管道脉冲涡流扫查服务": "石油领域的招投标", "玉门油田环庆分公司2024年度计量器具检定校准检测项目": "石油领域的招投标", "甘油代储代销": "石油领域的招投标", "秦皇岛-丰南沿海输气管道工程PC总承包五标段（收发球筒及分离器）采购": "石油领域的招投标", "西南油气田分公司勘探事业部2024年度溢流监测预警系统技术服务": "石油领域的招投标", "西南管道公司2024年B型套筒物资采购项目": "石油领域的招投标", "西部钻探2024年井下作业公司前置蓄能压裂注气服务(X09-023)定商项目": "石油领域的招投标", "辽阳石化分公司10万吨年尼龙66项目施工总承包": "石油领域的招投标", "长城钻探昆山公司2024年度国际货代": "石油领域的招投标", "长庆油田分公司第三采油厂2024年油井智能分采工艺技术改进研究与应用服务项目": "石油领域的招投标", "长庆油田分公司第三采油厂2024年盐67区二氧化碳驱注入管柱镍钨镀层防腐加工服务项目": "石油领域的招投标", "长庆油田分公司第二采油厂2024年环保管家技术服务项目": "石油领域的招投标", "长庆油田分公司第十二采油厂2024年轻烃生产辅助业务（二标段三次）": "石油领域的招投标", "长庆油田页岩油开发分公司2024年活动洗井车（带压密闭）洗井技术服务（二次）": "石油领域的招投标", "风城油田夏子街转油站完善工程": "石油领域的招投标", "中华人民共和国中国人民银行法": "法规", "中华人民共和国企业国有资产法": "法规", "中华人民共和国企业破产法": "法规", "中华人民共和国会计法": "法规", "中华人民共和国保险法": "法规", "中华人民共和国反不正当竞争法": "法规", "中华人民共和国商业银行法": "法规", "中华人民共和国审计法": "法规", "中华人民共和国招标投标法": "法规", "中华人民共和国招标投标法实施条例": "法规", "中华人民共和国社会保险法": "法规", "中华人民共和国预算法": "法规", "中华人民共和国预算法实施条例": "法规", "平安产险交通工具意外伤害保险（互联网版）条款": "法规", "平安产险女性安康特定疾病保险（互联网版）条款": "法规", "平安产险学生意外伤害保险（互联网版）条款": "法规", "平安产险意外伤害保险（互联网版）条款": "法规", "平安产险老年人医疗费用保险（互联网版）条款": "法规", "平安产险附加家庭财产服务保险条款": "法规", "平安产险附加特定药品费用医疗保险（B款）条款": "法规", "平安宠物传染病医疗费用损失保险条款": "法规", "平安家庭财产保险（家庭版）条款": "法规", "平安银行卡盗刷保险条款": "法规", "政府购买服务管理办法": "法规", "政府采购货物和服务招标投标管理办法": "法规", "政府采购非招标采购方式管理办法": "法规", "期货交易管理条例": "法规", "私募投资基金监督管理暂行办法": "法规", "证券公司监督管理条例": "法规", "金融违法行为处罚办法": "法规", "2024年1-4月中国房地产企业销售业绩排行榜": "行业研报", "2024年4月博彩数据跟踪点评：4月博收超出预期，黄金周客流略低于预期": "行业研报", "交通运输仓储行业研究：马士基继续暂停船只通过红海，2024年航空客流力争6.9亿人次": "行业研报", "传媒互联网周报：OpenAI、小冰等国多模态AI应用落地，“数据要素X”三年行动计划正式印发": "行业研报", "传媒互联网行业周报：五一档票房破15亿，假期大涨维持看好恒生科技": "行业研报", "传媒行业周报：国产大模型多模态方向发展，积极探索商业化": "行业研报", "公用事业—电力天然气周报：电力市场监管步入正轨，4月天然气进口量同比增长13.7%": "行业研报", "化工行业周报：海外天然气价格上涨，纯碱、草甘膦价格上涨": "行业研报", "化工行业周报：细分板块关注度提升，产品价格延续涨势": "行业研报", "医药制造行业2024年度行业分析": "行业研报", "商用车：4月重卡销量略超预期，出口+天然气持续高增": "行业研报", "家电行业周报：线上扫地机表现亮眼，洗地机维持增长": "行业研报", "影视行业行业动态：后续电影档期展望，复苏已变新成长": "行业研报", "我国特高压建设提速，柔性直流输电市场潜力巨大": "行业研报", "房地产行业周报（2024年第18周）：政治局会议新表述注入信心，地产股政策博弈价值凸显": "行业研报", "房地产：2024年5月上海市住宅价格发布": "行业研报", "房地产：关于深圳城中村改造新政意见稿的点评-深圳城中村改造提速推进，新政意见稿最大变化在于向一二级分离倾斜": "行业研报", "有色金属行业月度点评：金属价格全面上行，有色板块表现强势": "行业研报", "汽车板块2023年年报&2024年一季报总结：2023年汽车各子板块业绩均改善，2024年Q1客车业绩超预期": "行业研报", "消费行业五一消费观察：常态化复苏，结构性特征延续": "行业研报", "游戏Ⅱ行业深度报告：游戏行业2023年及2024Q1业绩综述：产品周期驱动业绩增长，全方位拥抱AI技术变革": "行业研报", "煤炭行业周报（1月第2周）：动力煤价格反弹，1月煤价将维持高位": "行业研报", "电力行业3月月报：中电联预计上半年用电量增速有望超8%，现货电价环比逐步企稳": "行业研报", "电力设备及新能源行业点评：Solaredge一季度续亏，欧美户用逆变器市场较弱": "行业研报", "电子行业2023年报及2024一季报综述：周期拐点已现，AI动能持续": "行业研报", "电子行业周观点：AI存储需求高景气，晶圆大厂切入先进封装赛道": "行业研报", "电子行业：2024Q1 PCB需求回暖，AI仍是主要驱动力": "行业研报", "石化周报：地缘风险溢价回落，关注OPEC+减产情况": "行业研报", "社会服务：五一黄金周点评-国内出行韧性凸显，多重利好推动出境游增长靓丽": "行业研报", "计算机行业周报：特斯拉通过国家汽车数据安全要求，自动驾驶板块有望受益": "行业研报", "计算机行业周观点：《“数据要素×”三年行动计划（2024—2026年）》印发，自定义Chat GPT商店即将上线": "行业研报", "计算机：交通数字化升级政策出台，带来千亿增量市场": "行业研报", "通信行业周观点：工信部等十一部门联合印发《关于开展“信号升格”专项行动的通知》": "行业研报", "部分长丝装置停车改造化工品价差继续扩大": "行业研报", "金属行业5月投资策略：继续看好有色上游机会": "行业研报", "钙钛矿行业深度报告：产研并进，降本提效，共赴星辰大海": "行业研报", "钢铁行业跟踪周报：需求弱势叠加成本上行，预计钢价震荡运行": "行业研报", "银行理财产品周报": "行业研报", "银行行业：金融支持住房租赁发展，有望成为新业务增长点": "行业研报", "非银金融行业事件点评报告：股息税减免预期催化大涨，关注政策α和复苏周期β机遇": "行业研报"}

		self.error_types = ["数值缺失错误", "语句缺失错误" ,
							"格式错误", "时间信息非法", "数值单位错误", 
							"主体不一致", "冗余语句", "时间矛盾错误", "计算错误", "语义逻辑矛盾", "数值不一致错误"]
		self.target_errors = ['数值缺失错误', '语句缺失错误', '格式错误', "时间信息非法", "数值单位错误", "计算错误"]
		  #["数值单位错误", "计算错误", "格式错误", "数值缺失错误"]

		self.prompt = """请你根据给定的文本片段，插入指定类型的错误数据，要求尽可能少的改动原始片段的内容，且插入的错误利用常识或通过上下文即可判断。若确实无法插入该类错误，请返回[]。
		1. 【错误类型定义】：
		2. 【输出要求】：只需输出被插入错误的语句(对)。
		3. 【示例】"""	
		self.defs = {"数值缺失错误": "数值缺失错误是指在文本中出现了应当明确提供的数值信息（如数量、金额、比例、时间、网址等），但该数值被遗漏，导致语义不完整或信息表达不清的情况。",
			"语句缺失错误": "语句缺失错误是指文段中存在结构性或内容性遗漏，尤其是在列举、排比、并列句等句式中，缺失了应有的语句、成分或内容项，导致语义不完整或表述突兀。",
			"格式错误":"格式错误是指文档中出现的数据项在表现形式上不符合预先设定或常用的格式规则。常见问题包括：网址不全，邮编不是6位数，邮箱少@，联系电话位数不对。",
			"时间信息非法": "时间信息非法是指文档中出现了不符合现实或常规表达的时间描述。常见问题包括：日期中的“日”超过该月的最大天数，月份超过12月，小时数不在0-23的范围内，年份不符合上下文语境。",
			"数值单位错误":"数值单位错误是指文档中数值所对应的单位使用不当。常见问题包括：单位选择不合常理（如面积使用平方分米，观影人数使用“人次”但具体数量为小数）；单位与数值语义不匹配（如礼堂改造项目的预算金额为200.00元）。",
			"主体不一致":"主体不一致是指文档中提及的主体（如人、组织、事件等）在不同部分的描述存在冲突或混淆。常见问题包括：不同主题被误认为同一主体；属性描述相互矛盾。",
			"冗余语句":"冗余语句是指文档中出现内容重复或表达重复的语句片段，这类冗余通常不会增加新的信息。常见问题包括：相同数据；措辞的重复。",
			"时间矛盾错误":"时间矛盾错误是指文档中出现了违反时间逻辑的描述。常见问题包括：逻辑时间顺序错误（如“开标早于投标截至时间”）;时间表达不合理（如“上午23点”）。",
			"计算错误":"计算错误是指文档中涉及数值计算、比例分配、增长率变动等方面的结果与其应有的计算逻辑不一致。常见问题包括：数值计算错误，比例之和大于100%，增减幅度计算错误等。",
			"语义逻辑矛盾":"语义逻辑矛盾是指文档内容在逻辑上存在明显冲突或自相矛盾的情况。常见问题包括：条件与结论不符；排名与所属范围不一致；表述本身存在矛盾或语义冲突。",
			"数值不一致错误":"数值不一致错误是指在同一文档中，本应保持一致的数值信息在不同位置出现不一致的情况，通常发生在标书编号、金额、数量、时间等关键字段。"
			}	
		self.error_examples = {"数值缺失错误": [["上周全国高速公路货车通行量为万辆"], ["联系方式："], ["年初至今出口额、出口量同比分别-4%。"]],
			"语句缺失错误": [["重点关注：1）AI+影视：芒果超媒；2）AI+IP（涉及版权、算力等）：凤凰传媒、山东出版；3）AI+游戏。"], ["价格方面，主要港口一级冶金焦价格2540元/吨，周环比下跌3.79%；二级冶金焦价格2340元/吨。"]],
			"格式错误": [["联系方式：10-"], ["邮编：132"]],
			"时间信息非法": [["4月31日"], ["129年发布"], ["晚上25时"]],
			"数值单位错误": [["昨日观影数量9.6人次"], ["建筑面积1590平方分米"], ["总金额：315.0000元"]],
			"主体不一致": [["项目名为：丙酮运输招标", "投标人需具备乙醇相关资质"], ["订立本保险合同时，采用保险人提供的格式条款的，保险人向投保人提供的投保单应当附格式条款，被保险人应当向投保人说明本保险合同的内容。"]],
			"冗余语句":[["建筑控制规模84979㎡，容积率2.8，控高60米，容积率2.8"]],
			"时间矛盾错误":[["上午13点"],["投标截至时间：2024年3月12日9点", "开标时间：2024年3月11日9点"]],
			"计算错误":[["预算金额：1155.845万元，第一次付款30%，第二次付款40%，第三次付款40%"], ["同比从+50.57%上升24.11pct至+84.68%"]],
			"语义逻辑矛盾":[["参与投标的供应商必须具备以下资格，任何一条不满足将不会被否决参与投标。"], ["在申万31个一级子行业中排名为第35位"], ["3．投标人资格要求","3.1 被列入失信人名单得企业。"]],
			"数值不一致错误": [["项目编号：GC-HGX240192","长庆油田分公司第三采油厂2024年油井智能分采工艺技术改进研究与应用服务项目GC-HGX240195"], ["总预算：315.6万元", "本项目的预算总金额为316.6万人民币"]]
			}
		self.icl_examples = {"数值缺失错误": [["...上周全国高速公路货车通行量为1000万辆...", "上周全国高速公路货车通行量为万辆"], ["...热泵单2023年13月份出口额、出口量分别为3.9亿、6.3亿台，同比分别-40%、-36%，年初至今出口额、出口量同比分别-4%,-6%。...", '热泵单2023年13月份出口额、出口量分别为3.9亿、6.3亿台，同比分别-40%、-36%，年初至今出口额、出口量同比分别-4%。']],
			"语句缺失错误": [["...重点关注：1）AI+影视：芒果超媒；2）AI+IP（涉及版权、算力等）：凤凰传媒、山东出版；3）AI+游戏：美国艺电。","重点关注：1）AI+影视：芒果超媒；2）AI+IP（涉及版权、算力等）：凤凰传媒、山东出版；3）AI+游戏：。"], ["...价格方面，主要港口一级冶金焦价格2540元/吨，周环比下跌3.79%；二级冶金焦价格2340元/吨，周环比下跌1.2%。...", "价格方面，主要港口一级冶金焦价格2540元/吨，周环比下跌3.79%；二级冶金焦价格2340元/吨。"]],
			"格式错误": [["...联系方式：0990-6847512", "联系方式：10-"], ["...邮编：100053", "邮编：132"]],
			"时间信息非法": [["...法人、法定代表人自2017年1月1日至投标截止时间无行贿犯罪；...", "法人、法定代表人自2017年13月1日至投标截止时间无行贿犯罪；"], ["...2023年第二季度以来，海外户用光伏和储能系统库存积压问题较为严重...", "223年第二季度以来，海外户用光伏和储能系统库存积压问题较为严重"], ["...投标截至时间：2024年3月12日9时", "投标截至时间：2024年3月12日29时"]],
			"数值单位错误": [["...昨日观影数量9.6万人次。...","昨日观影数量9.6人次"], ["...建筑面积1590平方米...","建筑面积1590平方分米"], ["...总金额：315.0000万元...","总金额：315.0000元"]],
			"主体不一致": [["项目名为：丙酮运输招标...投标人需具备丙酮相关资质", "投标人需具备乙醇相关资质"], ["...订立本保险合同时，采用保险人提供的格式条款的，保险人向投保人提供的投保单应当附格式条款，保险人应当向投保人说明本保险合同的内容。...", "订立本保险合同时，采用保险人提供的格式条款的，保险人向投保人提供的投保单应当附格式条款，被保险人应当向投保人说明本保险合同的内容。"]],
			"冗余语句":[["...建筑控制规模84979㎡，容积率2.8，控高60米。...", "建筑控制规模84979㎡，容积率2.8，控高60米，容积率2.8。"]],
			"时间矛盾错误":[["...购买时间：2024年05月11日至2024年05月17日23时59分止。...", "购买时间：2024年05月11日至2024年05月07日23时59分止。"],["...投标截至时间：2024年3月12日9点...开标时间：2024年3月12日9点", "投标截至时间：2024年3月12日9点", "开标时间：2024年3月11日9点"]],
			"计算错误":[["...预算金额：1155.845万元，第一次付款30%，第二次付款40%，第三次付款30%...", "第一次付款30%，第二次付款40%，第三次付款40%"], ["同比从+50.57%上升24.11pct至+74.68%", "同比从+50.57%上升24.11pct至+84.68%"]],
			"语义逻辑矛盾":[["...参与投标的供应商必须具备以下资格，任何一条不满足将会被否决参与投标。", "参与投标的供应商必须具备以下资格，任何一条不满足将不会被否决参与投标。"], ["...在申万31个一级子行业中排名为第3位", "在申万31个一级子行业中排名为第35位"], ["3．投标人资格要求...3.1 未被列入失信人名单得企业。", "3．投标人资格要求", "3.1 被列入失信人名单得企业。"]],
			"数值不一致错误": [["项目编号：GC-HGX240192...长庆油田分公司第三采油厂2024年油井智能分采工艺技术改进研究与应用服务项目（GC-HGX240192）", "项目编号：GC-HGX240192", "长庆油田分公司第三采油厂2024年油井智能分采工艺技术改进研究与应用服务项目（GC-HGX240195）"], ["总预算：315.6万元...本项目的预算总金额为315.6万人民币", "总预算：315.6万元", "本项目的预算总金额为316.6万人民币"]]
			}	
	
	def read_docx(self, doc_file):
		doc = Document(doc_file)
		paragraphs = []

		name = os.path.splitext(doc_file)[0]
		name = name.split('/')[-1]

		first_para = doc.paragraphs[0].text
		overlap = len(set(first_para) & set(name))*1.0 / len(set(first_para) | set(name))

		if overlap < 0.5:
			paragraphs.append(first_para)

		for para in doc.paragraphs[1:]:
			text = para.text.strip().replace(' ', ' ').replace('　', '').replace(' ', ' ').replace('    ', '    ').replace(' ', ' ')
			if text:
				paragraphs.append(text)

		self.doc_lens.append(len('\n'.join(paragraphs)))
		return {
			"type": self.doc_types[name],
			"name": name,
			"paras": paragraphs
			}

	def filter_sents(self, sents):
		#过滤一部分肯定不能生成插入指定错误的文本
		pa = r'[^a-z]+[0-9]+[^a-z]+'	
		pa1 = r'(邮箱|联系方式|网址|电话).*?[0-9a-z]+'
		pa2 = r'([0-9]{0,2}月[0-9]{0,2}日)|([0-9]+日[0-9]+时)'
		pa3 = r'(金额|预算|面积|数量|人|单价|元|辆|载重).*[0-9]+'
		p1 = r'\n([0-9]\.[0-9]\.[0-9]\.*)|([0-9]\.[0-9])|([0-9]\.*)'
		cand_sents = {"格式错误":[], 
					"时间信息非法":[],
					"数值单位错误":[],
					"语句缺失错误":[],
					"数值缺失错误":[],
					"时间矛盾错误":[],
					"计算错误":[],
					"语义逻辑矛盾":[]
					}
		print_flag = False
		for s in sents:
			if '招标公告中未尽事宜或与招标文件不符之处，以招标文件为准' in s or '投标文件的递交' in s:
				print_flag = True
			new_s = re.sub(p1, '', s)
			if re.search(pa1, new_s):
				cand_sents['格式错误'].append(s)
				cand_sents['数值缺失错误'].append(s)
				if print_flag:
					print('格式错误', re.search(pa1, new_s))
				new_s = re.sub(pa1, '', new_s)	
			if re.search(pa2, new_s):
				cand_sents['时间信息非法'].append(s)
				if print_flag:
					print('时间信息非法', re.search(pa2, new_s))
			if len(re.findall(pa2, new_s)) >= 2:
				cand_sents['时间矛盾错误'].append(s)
				if print_flag:
					print('时间矛盾错误', re.search(pa2, new_s))
			if re.search(pa3, new_s):
				cand_sents['数值单位错误'].append(s)
				cand_sents['数值缺失错误'].append(s)
				if print_flag:
					print('数值单位错误', re.search(pa3, new_s))

			if re.search(r'[0-9]+', s):
				cand_sents['数值单位错误'].append(s)	
				
			if ('；' in new_s) or (';' in new_s):
				cand_sents['语句缺失错误'].append(s)

			if s not in cand_sents['数值缺失错误']:
				if re.search(pa, new_s):
					cand_sents['数值缺失错误'].append(s)
					if print_flag:
						print('数值缺失错误', re.search(pa, new_s))

			if s not in cand_sents['计算错误']:
				if len(re.findall(r'[0-9\.]+', s)) >= 3:
					cand_sents['计算错误'].append(s)

		return cand_sents

	def generate_error4single_doc(self, doc_info):
		#往一个doc中插入错误
		doc = '\n'.join(doc_info['paras'])
		sents = doc.split('。')

		if len(doc) > 1500:
			cand_sents = self.filter_sents(sents)

		prompts = []
		for error in self.target_errors:   #按照错误类型插入错误
			examples = '\n'
			for x in self.icl_examples[error]:
				origin = x[0]
				new = x[1:]
				examples += '输入：'+ origin+'\n输出：' + json.dumps(new, ensure_ascii=False)+'\n'
			
			prompt1 = self.prompt.replace('	', '').replace('【错误类型定义】：', '【错误类型定义】：'+self.defs[error]).replace('【示例】', '【示例】'+examples)

			if len(doc) < 1500:  #少于1500个字，直接把整个文档输给大模型插入错误		
				prompt1 = prompt1+'--------------------\n输入：'+ doc + '\n输出：'
				prompts.append(prompt1)
				# pass
			else:
				sents = cand_sents.get(error, [])
				for x in sents:
					prompt2 = prompt1 +'--------------------\n输入：'+ x + '\n输出：'
					prompts.append(prompt2)

		self.total_sents += len(prompts)			
		print('一共可以尝试插入错误的候选句子有: ', len(prompts))
		
		#调用大模型插入错误
		inject_res = []	
		step = 32
		for i in range(0, len(prompts), step):
			inject_res.extend(self.llm_api.get_multi_response(prompts[i:i+step]))    #一个问题会产生多个答案，[query, [ans1, ans2]]
		
		# inject_res.extend(self.llm_api.get_response([prompts[0]]))

		# print(len(inject_res))
		# print(inject_res)
		return prompts, inject_res	
	
	def get_errors(self, initial_res, initial_docs):
		#从大模型的返回结果中抽取出插入的错误
		#initial_res: {"name":[[prompt, [ress]]]}
		#initial_docs: {"name":{"name":"", "type":"", "paras":[]}}
		#final_errors: {"name", "error_type": [correct, [errors]]}
		pattern = r'\["([^\]]*?)"\]'    #只抽取[]里面的内容

		final_errors = {}
		for name, info in initial_res.items():
			final_errors[name] = {}
			for prompt, temp_errors in info:
				origin = prompt.split('--------------------\n输入：')[-1].split('\n输出：')[0]
				error_type = prompt.split('【错误类型定义】：')[-1].split('是指')[0]
				errors = []
				for temp_x in temp_errors:
					try:
						temp_res = re.findall(pattern, temp_x)
						if not temp_res:
							continue

						x = temp_res[-1]
						if '","' in x:
							x = x.split('","')
						else:
							x = [x]
						flag = False
						if len(x) == 1:
							for y in errors:
								if x[0] in y[0]:
									flag = True
									continue
							if not flag:
								errors.append(x)
						elif len(x) == 2:
							for y in errors:
								if len(y) == 1:
									if (x[0] in y[0]) and (x[1] in y[0]):
										flag = True
										continue
								else:
									if ((x[0] in y[0]) and (x[1] in y[0])) or ((x[0] in y[1]) and (x[1] in y[1])):
										flag = True
										continue
									elif ((x[0] in y[0]) and (x[1] in y[1])) or ((x[0] in y[1]) and (x[1] in y[0])):
										flag = True
										continue
							if not flag:
								errors.append(x)
						else:
							print('未考虑到的情况', temp_x, temp_res)
					except:
						traceback.print_exc()
						print('处理大模型生成的错误出错:', temp_x, temp_res)
						# exit()
				if errors:
					if error_type not in final_errors[name]:
						final_errors[name][error_type] = []
					final_errors[name][error_type].append([origin, errors])
			# break
		return final_errors

	def generate_inject_errors(self, errors):
		#errors:{"name", "error_type": [[correct, [errors]]}
		#将生成的错误插入到原始文档中
		# 输入：
		# 错误类型：语句缺失错误
		# 错误语句：处转让、分包项目金额5‰以上10‰以下的罚款；并处没收违法所得；
		# 原始文本：第七十六条中标人将中标项目转让给他人的，将中标项目肢解后分别转让给他人的，违反招标投标法和本条例规定将中标项目的部分主体、关键性工作分包给他人的，或者分包人再次分包的，转让、分包无效，处转让、分包项目金额5‰以上10‰以下的罚款；有违法所得的，并处没收违法所得；可以责令停业整顿；情节严重的，由工商行政管理机关吊销营业执照
		# 输出：
		# 第七十六条中标人将中标项目转让给他人的，将中标项目肢解后分别转让给他人的，违反招标投标法和本条例规定将中标项目的部分主体、关键性工作分包给他人的，或者分包人再次分包的，转让、分包无效，处转让、分包项目金额5‰以上10‰以下的罚款；并处没收违法所得；可以责令停业整顿；情节严重的，由工商行政管理机关吊销营业执照。
		
		prompt = """给定原始文本和人工构造的错误语句，请你精确定位错误语句在原始文本中对应的位置，执行语句替换操作，从而生成一段带错文本。要求：每次只会提供一个错误语句，如提供的是包含两条语句的列表[]，只需关注第二条语句的对应位置，并进行替换操作。请保持原始文本结构的完整性，仅替换目标语句。禁止修改非目标区域的任何文本内容。
		操作步骤：
		1. 通过字符串对比的方式，先找到错误语句在原始文本中的开始位置；
		2. 找到错误语句在原始文本中的结束位置；如果结束位置已经是原始文本的最后一句，则结束位置调整为原始文本的最后一个字符。
		3. 利用错误语句替换原始文本中开始位置到结束位置的内容，并输出得到的文本。
		【示例1】
		输入：
		错误类型：语句缺失错误
		错误语句：[\"投标人应在5.2规定的投标截止时间前通过“中国石油电子招标投标交易平台”递交电子投标文件；（为避免受网速及网站技术支持工作时间的影响，建议于投标截止时间前24小时完成网上电子投标文件的递交）投标截止时间前未被系统成功传送的电子投标文件将。\"]
		原始文本：投标人应在5.2规定的投标截止时间前通过“中国石油电子招标投标交易平台”递交电子投标文件；（为避免受网速及网站技术支持工作时间的影响，建议于投标截止时间前24小时完成网上电子投标文件的递交）投标截止时间前未被系统成功传送的电子投标文件将不被接受，视为主动撤回投标文件
		输出：
		投标人应在5.2规定的投标截止时间前通过“中国石油电子招标投标交易平台”递交电子投标文件；（为避免受网速及网站技术支持工作时间的影响，建议于投标截止时间前24小时完成网上电子投标文件的递交）投标截止时间前未被系统成功传送的电子投标文件将。
		---------
		【示例2】
		输入：
		错误类型：语句缺失错误
		错误语句：(一)；(二)不按照规定确定中标人；
		原始文本：第七十三条依法必须进行招标的项目的招标人有下列情形之一的，由有关行政监督部门责令改正，可以处中标项目金额10‰以下的罚款；给他人造成损失的，依法承担赔偿责任；对单位直接负责的主管人员和其他直接责任人员依法给予处分：\n(一)无正当理由不发出中标通知书；\n(二)不按照规定确定中标人；\n(三)中标通知书发出后无正当理由改变中标结果；\n(四)无正当理由不与中标人订立合同；\n(五)在订立合同时向中标人提出附加条件
		输出：
		第七十三条依法必须进行招标的项目的招标人有下列情形之一的，由有关行政监督部门责令改正，可以处中标项目金额10‰以下的罚款；给他人造成损失的，依法承担赔偿责任；对单位直接负责的主管人员和其他直接责任人员依法给予处分：\n(一)；(二)不按照规定确定中标人；\n(三)中标通知书发出后无正当理由改变中标结果；\n(四)无正当理由不与中标人订立合同；\n(五)在订立合同时向中标人提出附加条件
		---------
		【示例3】
		输入：
		错误类型：数值缺失错误
		错误语句：[\"同比近乎持平。\", \"同比减少了15.2亿元（截至2024年5月5日14时）。\"]
		原始文本：2024年五一档票房超15亿元，同比近乎持平。今年五一档从5月1日至5月5日，档期票房为15.2亿元，同比近乎持平（截至2024年5月5日14时）。
		输出：
		2024年五一档票房超15亿元，同比近乎持平。今年五一档从5月1日至5月5日，档期票房为15.2亿元，同比减少了15.2亿元（截至2024年5月5日14时）。
		---------
		下面请你完成上述错误语句和原始文本对的替换。
		输入：
		错误类型：
		"""
		prompt = prompt.replace('	', '')
		prompts = []
		inject_res = []

		for name, info in errors.items():
			for error_type, e_info in info.items():
				for (x_cor, x_errors) in e_info:
					for y in x_errors:
						prompt1 = prompt + error_type +'\n错误语句：' + json.dumps(y, ensure_ascii=False) +'\n原始文本：'+x_cor+'\n输出：'
						prompts.append(prompt1)
						inject_res.append([name, error_type, x_cor, y])

		step = 16
		for i in range(0, len(prompts), step):
			temp_res = self.llm_api.get_multi_response(prompts[i:i+step])  #[query, [ans1, ans2]]
			# print("将生成的错误插入到原始文档中的中间结果temp_res:",)
			# print(json.dumps(temp_res, ensure_ascii=False))
			for j, (prompt, info) in enumerate(temp_res):
				temp_res = []
				for y in info:
					# y = y.split('\n')[0]
					if y not in temp_res:
						temp_res.append(y)
				inject_res[i+j].append(temp_res)

		return inject_res					

	def generate_instruction_data(self, errors, initial_docs, a_type='yn'):
		#initial_docs: {"name":{"name":"", "type":"", "paras":[]}}
		#errors: [[name, error_type, x_cor, [error], [[x_err]]
		# "中亚天然气管道（霍尔果斯）有限公司 2024年流量计使用中测试项目（二次）",
        # "语句缺失错误",
        # "\n②购买招标文件时投标商需切换至《中国石油招标中心数字化管理平台》缴费，其登陆地址：http://www2.cnpcbidding.com/#/login；具体流程及操作步骤详见《中国石油招标中心投标商操作手册V1.1》",
        # [
        #     "购买招标文件时投标商需切换至《中国石油招标中心数字化管理平台》缴费，其登陆地址：http://www2.cnpcbidding.com/#/login；具体流程及操作步骤。"
        # ],
        # [
        #     "②购买招标文件时投标商需切换至《中国石油招标中心数字化管理平台》缴费，其登陆地址：http://www2.cnpcbidding.com/#/login；具体流程及操作步骤。"
        # ]
		if a_type =='yn':
			prompt = """你是一名文档质量检查者，请依据给出的【错误类型定义】判断下述文本中是否存在该类错误，若存在该类错误，请回答“是”，否则请回答否”。并在第二行给出你的判断理由。
			【错误类型定义】：
			"""
		elif a_type == 'select':
			prompt = """你是一名金融文档质量检查者，金融文档中出现的错误大致可划分成以下十一大类：
			（1）数值缺失错误是指在文本中出现了应当明确提供的数值信息（如数量、金额、比例、时间、网址等），但该数值被遗漏，导致语义不完整或信息表达不清的情况。
			（2）语句缺失错误是指文段中存在结构性或内容性遗漏，尤其是在列举、排比、并列句等句式中，缺失了应有的语句、成分或内容项，导致语义不完整或表述突兀。
			（3）格式错误是指文档中出现的数据项在表现形式上不符合预先设定或常用的格式规则。常见问题包括：网址格式不对，邮编不是6位数，邮箱少@，联系电话位数不对。
			（4）时间信息非法是指文档中出现了不符合现实或常规表达的时间描述。常见问题包括：日期中的“日”超过该月的最大天数，月份超过12月，小时数不在0-23的范围内，年份不符合上下文语境。
			（5）数值单位错误是指文档中数值所对应的单位使用不。常见问题包括：单位选择不合常理（如面积使用平方分米，观影人数使用“人次”但具体数量为小数）；单位与数值语义不匹配（如礼堂改造项目的预算金额为200.00元）。
			（6）主体不一致是指文档中提及的主体（如人、组织、事件等）在不同部分的描述存在冲突或混淆。常见问题包括：不同主题被误认为同一主体；属性描述相互矛盾。
			（7）冗余语句是指文档中出现内容重复或表达重复的语句片段，这类冗余通常不会增加新的信息。常见问题包括：相同数据；措辞的重复。",
			（8）时间矛盾错误是指时间矛盾错误是指文档中出现了违反时间逻辑的描述。常见问题包括：逻辑时间顺序错误（如“开标早于投标截至时间”）;时间表达不合理（如“上午23点”）。
			（9）计算错误是指文档中涉及数值计算、比例分配、增长率变动等方面的结果与其应有的计算逻辑不一致。常见问题包括：数值计算错误，比例之和大于100%，增减幅度计算错误等。
			（10）语义逻辑矛盾是指文档内容在逻辑上存在明显冲突或自相矛盾的情况。常见问题包括：条件与结论不符；排名与所属范围不一致；表述本身存在矛盾或语义冲突。
			（11）数值不一致错误是指在同一文档中，本应保持一致的数值信息在不同位置出现不一致的情况，通常发生在标书编号、金额、数量、时间等关键字段。
			====
			现在请你跟据以上各类错误类型的定义判断下述文本中是否存在这十一类错误，若存在错误，请以json列表形式返回存在的错误类型和对应的错误句子（对），否则返回空json{}。具体json示例如下：
			```
			{
				"error": "数值缺失错误",
				"sents":["联系人及电话：李三。"]
			}
			```
			并请在第二行给出你的分析过程。请注意每个文本中最多只会存在一种错误，请你选择最合适的一种错误类型。
		"""
		else:
			prompt = "你是一名文档质量检查者，请依据错误类型定义判断下述文本中是否存在【【错误类型】】，若存在该类错误，请以嵌套列表形式返回存在该类错误的句子（对），否则返回空列表[]，并在第二行给出你分析错误语句的过程。\n【错误类型定义】："
		input1 = "【句子】："

		prompt = prompt.replace('	', '')
		input1 = input1.replace('	', '')
		data = []
		# print("大模型生成的训练数据集errors", json.dumps(errors, ensure_ascii=False))
		for info in errors:
			try:
				name, error_type, x_cor, error, x_err = info
			except:
				traceback.print_exc()
				print(info)
				exit()
			# x_err = x_err[-1]

			#判断下生成的新文本是否符合要求，
			err_text = error[-1].replace('\n', '')
			if len(x_err) > 1:  #按理说插入错误后只有一种结果，如果有多种，说明插入有问题，直接跳过
				if len(x_err[0]) < len(x_err[1]):
					short = x_err[0] 
					long_ = x_err[1]
				else :
					short = x_err[1]
					long_ = x_err[0]
				if long_.startswith(short) and len(long_.replace(short, '')) < 3:   #可能默认价格标点符号，这种情况也允许存在
					x_err = long_
				else:
					print('error: ',error_type, '个数大于1，说明插入存在问题')
					print('x_err：')
					print(json.dumps(x_err, ensure_ascii=False))
					continue
			else:
				x_err = x_err[0]

			idx = x_err.replace('\n', '').find(err_text)
			if idx == -1:
				print('error: ',error_type, '生成的新文本并没有插入错误语句error: ', error)
				print('插入后的文本 x_err：', x_err)
				# print(json.dumps(x_err, ensure_ascii=False))
				continue
			#然后判断其他内容是否被修改
			pre_text = x_err.replace('\n', '')[:idx]
			if pre_text not in x_cor.replace('\n', '') :
				print('error: ',error_type, '生成的新文本修改了错误语句之外的内容，前半段', pre_text, '\n-----------\n原始句子：', x_cor)
				continue
			suf_text = x_err.replace('\n', '')[idx+len(err_text)+2:]
			if suf_text.rstrip('。') not in x_cor.replace('\n', ''):
				print('error: ',error_type, '生成的新文本修改了错误语句之外的内容，后半段', suf_text,'\n-----------\n原始句子：',  x_cor)
				continue
				
			# x_err = json.dumps(x_err, ensure_ascii=False)
			if a_type == 'select':
				data.append({"instruction": prompt, "input": x_err, "output": json.dumps({"error": error_type, "sents": error}, ensure_ascii=False), "origin":x_cor, "doc_name":name})
				#删除原始文本中包含该错误的段落，遗留的用于生成无错示例
				flag = False if initial_docs[name]['paras'] else True
				for i, para in enumerate(initial_docs[name]['paras']):
					if (x_cor in para) or (para in x_cor):
						del initial_docs[name]['paras'][i]
						flag = True
						break
				if not flag:
					print("x_cor没有在原始文档中找到", name, x_cor)
					print("\n".join(initial_docs[name]['paras']))
				continue
			else:
				prompt1 = prompt.replace('【错误类型】', error_type).replace('【错误类型定义】：', '【错误类型定义】：'+self.defs[error_type])
				input11 = input1.replace('【句子】：', x_err)
				data.append({"instruction": prompt1, "input": input11, "output": json.dumps([error], ensure_ascii=False), "output1":"是", "origin":x_cor, "doc_name":name})

			error_types = ["数值缺失错误", "语句缺失错误",
							"格式错误", "时间信息非法", "数值单位错误", "计算错误",]
			try:
				error_types.remove(error_type)
			except:
				traceback.print_exc()
				print("类型错误",error_type)

			idx = random.randint(0, len(error_types)-1)
			temp_type = error_types[idx]
			try:
				prompt1 = prompt.replace('	', '').replace('【错误类型】', temp_type).replace('【错误类型定义】：', '【错误类型定义】：'+self.defs[temp_type])
			except:
				traceback.print_exc()
				print(temp_type, error_type)
				continue
			data.append({"instruction": prompt1, "input":input11, "output": json.dumps([], ensure_ascii=False), "output1":"否", "origin":x_cor, "doc_name":name})
		
		print('一开始提供的数据:', len(errors), '生成的微调数据:', len(data))
		
		if a_type == 'select':  #select要单独生成负样例
			for name, info in initial_docs[name]:
				for p in info['paras']:
					data.append({"instruction": prompt, "input": p, "output": "{}", "doc_name":name})

		return data

	def generate_reason(self, data):
		#给定包含错误的文本，让大模型抽取其中的错误语句，并给出理由
		#data :[{"instruction":, "input":,"output":,"origin":,"doc_name":}]

		prompts = []
		for x in data:
			input1 = x['input']
			prompts.append(x['instruction']+'-----\n'+input1)

		step = 20
		for i in range(0, len(prompts), step):
			temp_res = self.llm_api.get_multi_response(prompts[i:i+step])
			for j, (prompt, info) in enumerate(temp_res):
				data[i+j]['reason'] = info

		return data			

	def filter_finetune_data(self, data, a_type='yn'):
		#删除大模型的返回结果和标准答案不一致的情况生成微调的数据
		#data :[{"instruction", "input","output":[],"output1":"是","origin","doc_name", "reason"}]
		#a_type表示目标任务是判断还是抽取
		final_data = []
		if a_type == 'yn':
			for info in data:
				reasons = info['reason']
				for x in reasons:
					y = x.split('\n')
					if y[0].startswith(info['output1']):
						final_data.append({"instruction": info["instruction"], 
											"input":info['input'],
											"output":x,
											"origin": info['origin'],
											"doc_name": info['doc_name']})
		elif a_type == 'select':
			prompt = """你是一名金融文档质量检查者，金融文档中出现的错误大致可划分成以下十大类：
			（1）数值缺失错误;
			（2）语句缺失错误;
			（3）格式错误;
			（4）时间信息非法;
			（5）数值单位错误;
			（6）主体不一致;
			（7）冗余语句;
			（8）时间矛盾错误；
			（9）计算错误;
			（10）语义逻辑矛盾;
			（11）数值不一致错误;
			====
			现在请你跟据以上各类错误类型的定义判断下述文本中是否存在这十一类错误，若存在错误，请以json列表形式返回存在的错误类型和对应的错误句子（对），否则返回空json"{}"。具体json示例如下：
			```
			{
				"error": "数值缺失错误",
				"sents":["联系人及电话：李三。"]
			}
			```
			并请在第二行给出你的分析过程。请注意每个文本中最多只会存在一种错误，请你选择最合适的一种错误类型。
		"""
			pa = r'json[\n\s]*{.*?}'
			for info in data:
				reasons = info['reason']
				ground_th = json.load(info['output'])
				for x in reasons:
					res = re.findall(pa, x)
					if res:
						if len(res) > 1:    #找到了多个答案，删除
							continue
						elif info['output'] == "{}" and res[0] == "{}":
							final_data.append({"instruction": prompt, "input": info['input'], "output":"<|reason_start|>"+x.split('\n')[1]}+"<|reason_end|>"+'\n\n<|answer_start|>\n```json\n'+info['output']+'\n```\n<|answer_end|>)
						elif info['output'] != "{}" and res[0] != "{}":   #两个都有答案
							res = json.loads(res[0])
							if res['error'] != ground_th['error']:
								continue
							elif (res['sents'][0] in ground_th['sents'][0]) or (ground_th['sents'][0] in res['sents'][0])
								final_data.append({"instruction": prompt, "input": info['input'], "output":"<|reason_start|>"+x.split('\n')[1]}+"<|reason_end|>"+'\n\n<|answer_start|>\n```json\n'+info['output']+'\n```\n<|answer_end|>)
								
		else:
			old_prompt = """你是一名文档质量检查者，请依据给出的【错误类型定义】判断下述文本中是否存在该类错误，若存在该类错误，请回答“是”，否则请回答否”。并在第二行给出你的判断理由。"""
			prompt = "你是一名文档质量检查者，请依据错误类型定义判断下述文本中是否存在该类错误。请在第一行输出你分析错误语句的过程，并在第二行嵌套列表形式返回找到的错误句子（对）。"

			for info in data:
				reasons = info['reason']
				instruction = info['instruction']
				error_type = instruction.split('【错误类型定义】：')[-1].split('是指')[0]
				instruction = "你是一名文档质量检查者，请你判断下述文本中是否存在<|error_type_start|>"+error_type+"<|error_type_end|>。请在第一行输出分析过程，并在第二行以json格式返回找到的错误句子（对）。"
				for x in reasons:
					y = x.split('\n')
					if y[0].startswith(info['output1']):
						final_data.append({"instruction": instruction, 
											"input":info['input'].replace('输入：\n【句子】：', '').replace('\n输出：\n', ''),
											"output":"<|reason_start|>"+y[1]+"<|reason_end|>"+'\n\n<|answer_start|>\n```json\n'+json.dumps({"sentences": json.loads(info['output'])}, ensure_ascii=False)+'\n```\n<|answer_end|>',
											"origin": info['origin'],
											"doc_name": info['doc_name']})
		return final_data

	def test(self):
		prompt = "请你根据提供的几个示例，帮忙修改下每个错误类型的定义。\n【原始定义】：\n【示例】："
		for error, defi in self.defs.items():
			prompt1 = prompt.replace('【原始定义】：', '【原始定义】：'+defi).replace('【示例】：', '【示例】：'+json.dumps(self.examples[error], ensure_ascii=False))
			print(prompt1)

def main():
	injector = Inject()

	doc_path = './groundtruth/docs1'
	# doc_path = 'E:\\AFAC2024赛题\\赛题final\\隐藏榜'
	final_res = {}
	final_prompts = {}
	docs = {}
	mid_res_file  = 'new_errors_mid_res_check_2.json'
	mid_res = {}
	with open(mid_res_file, 'r', encoding='utf-8') as fr:
		mid_res = json.load(fr)

	# for f in os.listdir(doc_path):
	# 	file_name = doc_path+'/'+f
	# 	print("当前处理的文件:", f)
	# 	doc = injector.read_docx(file_name)
	# 	docs[doc['name']] = docs

	# 	temp_ps ,temp_res = injector.generate_error4single_doc(doc)

	# 	final_prompts[doc['name']] =  temp_ps
	# 	mid_res[doc['name']] = temp_res

	# 	break
	# print(injector.total_sents)
	# print(json.dumps(final_prompts, ensure_ascii=False))

	# with open(mid_res_file, 'w', encoding='utf-8') as fw:
	# 	json.dump(mid_res, fw, ensure_ascii=False, indent=4)


	# print('从大模型的返回结果中抽取出插入的错误')
	# all_errors = injector.get_errors(mid_res, docs)	

	# print("all_errors: 抽取出的错误语句")
	# print(json.dumps(all_errors, ensure_ascii=False))
	# print('将生成的错误插入到原始文档中')
	# data = injector.generate_inject_errors(all_errors)

	# with open('data_w_errors_2.json', 'w', encoding='utf-8') as fw:
	# 	json.dump(data, fw, ensure_ascii=False, indent=4)

	with open('data_w_errors_2.json', 'r', encoding='utf-8') as fr:
		data = json.load(fr)
	
	with open('data_w_errors.json', 'r', encoding='utf-8') as fr:
		data.extend(json.load(fr))

	data = injector.generate_instruction_data(data,a_type='select')
	with open('data4reason_3.json', 'w', encoding='utf-8') as fw:
		json.dump(data, fw, ensure_ascii=False, indent=4)

	# # with open('data4reason_2.json', 'r', encoding='utf-8') as fw:
	# # 	data = json.load(fr)

	data = injector.generate_reason(data)
	with open('data_w_reason_3.json', 'w', encoding='utf-8') as fw:
		json.dump(data, fw, ensure_ascii=False, indent=4)

	# with open('data_w_reason_2.json', 'r', encoding='utf-8') as fr:
	# 	data = json.load(fr)

	data = injector.filter_finetune_data(data,a_type='select')
	with open('data4sft_w_reason_all_3.json', 'w', encoding='utf-8') as fw:
		json.dump(data, fw, ensure_ascii=False, indent=4)

	# for p in [10, 50, 70, 80, 90, 95, 100]:
	# 	print(p, np.percentile(injector.doc_lens, p))	
	# 10 1306.2
	# 50 2723.0
	# 70 4393.499999999999
	# 80 7833.6
	# 90 9309.900000000001
	# 95 10262.049999999997
	# 100 12155.0

def merge():
	#将构建的数据集合并
	# resf1 = 'new_errors_mid_res_check_1.json'
	# resf2 = 'new_errors_mid_res_added_2.json'
	# with open(resf1, 'r', encoding='utf-8') as fr:
	# 	data1 = json.load(fr)
	# with open(resf2, 'r', encoding='utf-8') as fr:
	# 	data = json.load(fr)

	# for name, info in data1.items():
	# 	if name in data:
	# 		data1[name].extend(data[name])
	# 		del data[name]
	# if data:
	# 	for name , info in data:
	# 		data1[name] = info

	# print(json.dumps(data1, ensure_ascii=False))
	#平衡每个类别的训练数据
	errors = {}
	with open('data4sft_w_reason_all_1.json', 'r', encoding='utf-8') as fr:
		data = json.load(fr)
	for x in data: 
		error = x['instruction'].split('<|error_type_start|>')[-1].split('<|error_type_end|>')[0]
		ans = False if '{\"error_sentences\": []' in x['output'] else True
		if error not in errors:
			errors[error] = {'y':[], 'n':[]}
		if ans:
			errors[error]['y'].append(x)
		else:
			errors[error]['n'].append(x)

	new_data = []
	for x, info in errors.items():
		print(x, len(info['y']), len(info['n']))
		if len(info['y']) > 50:
			temp = random.sample(info['y'], 50)
			new_data.extend(temp)
		else:
			new_data.extend(info['y'])
		if len(info['n']) > 100:
			temp = random.sample(info['n'], 50)
			new_data.extend(temp)
		else:
			new_data.extend(info['n'])	
	print(json.dumps(new_data, ensure_ascii=False))


if __name__ == '__main__':
	main()
	# merge()